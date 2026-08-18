"""Word + Excel deliverables — structured for instructor review."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.formatting.rule import ColorScaleRule
from openpyxl.worksheet.hyperlink import Hyperlink

NAVY = RGBColor(0x0B, 0x4F, 0x3A)
BLUE = RGBColor(0x0F, 0x94, 0x88)
GOLD = RGBColor(0xB7, 0x79, 0x1F)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MUTED = RGBColor(0x4D, 0x5D, 0x6F)
GREEN = RGBColor(0x0F, 0x7B, 0x4C)
RED = RGBColor(0xC0, 0x39, 0x2B)


def _rtl_p(p) -> None:
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _font(run, size=11, bold=False, color=DARK, name="Calibri"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:cs"), "Tahoma")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def _add(doc, text, size=11, bold=False, color=DARK, rtl=True):
    p = doc.add_paragraph()
    if rtl:
        _rtl_p(p)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text or "")
    _font(run, size=size, bold=bold, color=color)
    return p


def _h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    _rtl_p(p)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
    return p


def _shade(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_rtl(cell) -> None:
    for p in cell.paragraphs:
        _rtl_p(p)


def _cell(cell, text, *, bold=False, color=DARK, size=10, fill=None, rtl=True):
    cell.text = ""
    p = cell.paragraphs[0]
    if rtl:
        _rtl_p(p)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text if text is not None else ""))
    _font(run, size=size, bold=bold, color=color)
    if fill:
        _shade(cell, fill)
    # tighter
    for pp in cell.paragraphs:
        pp.paragraph_format.space_after = Pt(2)
        pp.paragraph_format.space_before = Pt(2)


def _table(doc, headers, rows, col_w=None, rtl=True):
    cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        _cell(tbl.rows[0].cells[i], h, bold=True, color=WHITE, size=10, fill="0B4F3A", rtl=rtl)
    for r_i, row in enumerate(rows, 1):
        fill = "F4F7FB" if r_i % 2 == 0 else "FFFFFF"
        for c_i in range(cols):
            val = row[c_i] if c_i < len(row) else ""
            _cell(tbl.rows[r_i].cells[c_i], val, size=9, fill=fill, rtl=rtl)
    if col_w:
        for row in tbl.rows:
            for i, w in enumerate(col_w):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return tbl


def _kv_table(doc, pairs):
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(pairs):
        fill = "EEF2F6" if i % 2 == 0 else "FFFFFF"
        _cell(tbl.rows[i].cells[0], k, bold=True, color=NAVY, size=10, fill="E2F4F1")
        _cell(tbl.rows[i].cells[1], v, size=10, fill=fill)
        tbl.rows[i].cells[0].width = Cm(4.6)
        tbl.rows[i].cells[1].width = Cm(12.2)
    doc.add_paragraph()
    return tbl


def _page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _header_footer(doc, owner, product):
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.right_margin = Cm(1.6)
        section.left_margin = Cm(1.6)
        sectPr = section._sectPr
        bidi = OxmlElement("w:bidi")
        sectPr.append(bidi)

        hp = section.header.paragraphs[0]
        _rtl_p(hp)
        r = hp.add_run(f"تجارت‌یار  |  پرونده تصمیم‌گیری واردات  |  {owner}  |  {product}")
        _font(r, size=8, color=MUTED)

        fp = section.footer.paragraphs[0]
        _rtl_p(fp)
        r1 = fp.add_run("تجارت‌یار · طراحی و توسعه: ستایش جعفری · خروجی خودکار به‌تنهایی منبع قطعی نیست · صفحه ")
        _font(r1, size=8, color=MUTED)
        # PAGE field
        fld1 = OxmlElement("w:fldChar")
        fld1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        run = fp.add_run()
        run._r.append(fld1)
        run2 = fp.add_run()
        run2._r.append(instr)
        run3 = fp.add_run()
        run3._r.append(fld2)
        _font(run, size=8, color=MUTED)
        _font(run2, size=8, color=MUTED)
        _font(run3, size=8, color=MUTED)


def _slug(s: str) -> str:
    s = re.sub(r"[^\w\u0600-\u06FF\-]+", "_", s or "").strip("_")
    return s[:60] or "product"


def _safe_excel_value(value):
    """Prevent spreadsheet-formula injection from web/user supplied text."""
    if isinstance(value, str) and value[:1] in ("=", "+", "-", "@"):
        return "'" + value
    return value


def file_stems(dossier: dict) -> str:
    m = dossier["meta"]
    return f"{m['owner_fa']}_{_slug(m['product_fa'])}"


def build_docx(dossier: dict, out_dir: Path) -> Path:
    m = dossier["meta"]
    brief = dossier["brief"]
    market = dossier["market"]
    sourcing = dossier["sourcing"]
    scoring = dossier["scoring"]
    cards = dossier["cards"]
    rfq = dossier["rfq"]
    decision = dossier["decision"]
    sources = dossier["sources"]

    doc = Document()
    _header_footer(doc, m["owner_fa"], brief["name_fa"])

    # ----- COVER -----
    _add(doc, m.get("platform") or "تجارت‌یار", size=11, color=GOLD)
    _add(doc, m.get("project_title") or "پرونده تصمیم‌گیری واردات", size=26, bold=True, color=NAVY)
    _add(doc, "گزارش حرفه‌ای ارزیابی واردات و تأمین‌کنندگان", size=13, color=MUTED)
    doc.add_paragraph()
    _kv_table(
        doc,
        [
            ("تهیه‌کننده گزارش", f"{m['owner_fa']}  /  {m['owner_en']}"),
            ("سازمان / مجموعه", m.get("organization") or "—"),
            ("هدف پرونده", m.get("report_purpose") or "—"),
            ("محصول", f"{brief['name_fa']}"),
            ("Product", brief["name_en"]),
            ("تاریخ پرونده", m["generated_on"]),
            ("نسخه سامانه", m.get("agent_version") or "—"),
            ("طراحی و توسعه", f"{m.get('developer_fa','ستایش جعفری')} / {m.get('developer_en','Setayesh Jafari')}"),
            ("کد HS محتمل", brief.get("hs_primary") or "نیازمند راستی‌آزمایی"),
            ("تعداد Longlist", str(len(sourcing.get("longlist") or []))),
            ("کانال‌های کشف", " ، ".join(sourcing.get("channels_used") or []) or "—"),
            ("وضعیت تصمیم", decision.get("recommendation_status_fa") or "—"),
            ("گزینه اول مذاکره", decision.get("first_choice") or "انتخاب نشد"),
            ("گزینه دوم (پشتیبان)", decision.get("second_choice") or "انتخاب نشد"),
            ("کنترل کیفیت", (dossier.get("quality_assurance") or {}).get("status") or "—"),
            ("تعداد منابع پذیرفته‌شده", str(len(sources))),
        ],
    )
    _add(
        doc,
        "این پرونده برای تصمیم‌گیری است نه پیدا کردن ارزان‌ترین فروشنده. "
        "هر ادعای مهم باید به Source Log برگردد. وضعیت مجاز/مشروط/ممنوع، ثبت سفارش، "
        "استاندارد، تعرفه و محدودیت‌های ارزی باید پیش از هر تصمیم تجاری در سامانه رسمی کنترل شوند.",
        size=10,
        color=MUTED,
    )

    _h(doc, "چگونه این پرونده را بخوانید", 2)
    _table(
        doc,
        ["بخش", "چه چیزی را بررسی کنید"],
        [
            ["خلاصه اجرایی", "گزینه اول و دوم و سه دلیل اصلی"],
            ["مرحله ۱", "تعریف محصول، عبارات جستجو، دلیل انتخاب HS"],
            ["مرحله ۲", "شواهد واردات ایران + ریسک مقرراتی + پورتال رسمی"],
            ["مرحله ۳ و اکسل Longlist", "حداقل ۲۰ تأمین‌کننده از چند کانال + موارد حذف‌شده"],
            ["مرحله ۴ و اکسل امتیاز", "مدل ۱۰۰ امتیازی و دلیل هر نمره برای Top 5"],
            ["مرحله ۵", "کارت Due Diligence، Green/Red Flag، گواهی verifyنشده"],
            ["مرحله ۶ و فایل RFQ", "پرامپت، پیش‌نویس، شخصی‌سازی، سؤال فنی و اعتبارسنجی"],
            ["مرحله ۷", "مقایسه ۵ گزینه و کارهای باز"],
        ],
    )

    # ----- EXEC -----
    _page_break(doc)
    _h(doc, "خلاصه اجرایی", 1)
    _kv_table(
        doc,
        [
            ("گزینه اول", decision.get("first_choice") or "—"),
            ("دلایل گزینه اول", "\n".join(f"• {x}" for x in (decision.get("first_reasons") or [])[:4])),
            ("گزینه دوم", decision.get("second_choice") or "—"),
            ("دلایل گزینه دوم", "\n".join(f"• {x}" for x in (decision.get("second_reasons") or [])[:4])),
            ("HS", f"{brief.get('hs_primary') or '—'}  |  جایگزین: {', '.join(brief.get('hs_alternatives') or []) or '—'}"),
            ("بازار ایران", market.get("imported_statement") or ""),
        ],
    )
    _add(doc, decision.get("recommendation_status_fa") or "", size=11, bold=True, color=RED)
    _add(doc, decision.get("disclaimer") or "", size=10, color=GOLD)
    _h(doc, "کنترل کیفیت خودکار پرونده", 2)
    qa_rows = []
    for q in (dossier.get("quality_assurance") or {}).get("checks") or []:
        qa_rows.append(["قبول" if q.get("passed") else "ناقص/رد", q.get("check"), q.get("detail")])
    if qa_rows:
        _table(doc, ["وضعیت", "کنترل", "جزئیات"], qa_rows)

    # ----- TOOLS -----
    _page_break(doc)
    _h(doc, "نقشه ابزارها — در هر مرحله از چه چیزی و چطور استفاده شد", 1)
    _add(
        doc,
        "نام ابزار واقعی و روش استفاده ثبت شده است. site:domain فقط جستجوی هدفمند وب عمومی است و "
        "به معنای استفاده از API یا حساب آن سرویس نیست. نتیجه خام جستجو تا عبور از فیلترها منبع پذیرفته‌شده محسوب نمی‌شود.",
        size=10,
        color=MUTED,
    )
    tool_rows = []
    for r in dossier.get("tool_log") or []:
        qs = " | ".join((r.get("queries") or [])[:3])
        tool_rows.append([r.get("stage"), r.get("tool"), r.get("how") or r.get("method"), r.get("hits"), qs[:120]])
    if tool_rows:
        _table(doc, ["مرحله", "ابزار", "چگونه استفاده شد", "نتایج", "نمونه پرس‌وجو"], tool_rows)
    cat_rows = [[v["name"], v["role"], v["method"]] for v in (dossier.get("tool_catalog") or {}).values()]
    if cat_rows:
        _h(doc, "کاتالوگ ابزارهای این پرونده", 2)
        _table(doc, ["ابزار", "نقش در پروژه", "روش استفاده"], cat_rows)

    # ----- STAGE 1 -----
    _h(doc, "مرحله ۱ — Product Brief", 1)
    _kv_table(
        doc,
        [
            ("نام فارسی", brief["name_fa"]),
            ("نام انگلیسی", brief["name_en"]),
            ("کاربرد", brief["application"]),
            ("مشخصات فنی", brief["specs"]),
            ("گرید / مدل / ظرفیت", brief["grade_model"]),
            ("واحد خرید", brief["unit"]),
            ("مشتری هدف در ایران", brief["target_customer"]),
            ("مقدار تقریبی سفارش", brief["qty_hint"]),
            ("گروه محصول", brief.get("product_category_label") or "—"),
            ("کشورها/مبدأهای جستجو", "، ".join(brief.get("origin_strategy") or [])),
            ("ویژگی‌های مؤثر بر HS", " | ".join(brief.get("classification_attributes") or [])),
            ("کیفیت ورودی", f"{(brief.get('input_quality') or {}).get('score', 0)}٪ — {(brief.get('input_quality') or {}).get('status', '')}"),
            ("شرح Product Brief", (brief.get("description_web") or "—")[:700]),
        ],
    )
    _h(doc, "عبارات جستجوی انگلیسی (Sourcing)", 2)
    _table(
        doc,
        ["#", "عبارت"],
        [[str(i), p] for i, p in enumerate(brief.get("search_phrases") or [], 1)],
        rtl=False,
    )
    _h(doc, "کد HS و دلیل انتخاب", 2)
    hs_rows = [[brief.get("hs_primary") or "—", "اصلی", str((brief.get("hs_counts") or {}).get(brief.get("hs_primary") or "", "—"))]]
    for alt in brief.get("hs_alternatives") or []:
        hs_rows.append([alt, "جایگزین", str((brief.get("hs_counts") or {}).get(alt, "—"))])
    _table(doc, ["کد", "نقش", "تعداد مشاهده در منابع وب"], hs_rows)
    _add(doc, brief.get("hs_reason") or "", color=GOLD)
    for warning in (brief.get("input_quality") or {}).get("warnings") or []:
        _add(doc, "• " + warning, size=10, color=RED)
    _add(
        doc,
        "اگر چند کد دیده شد، اختلاف معمولاً به‌خاطر سطح تفکیک (۶ در برابر ۸/۱۰ رقمی)، گرید یا کاربرد نهایی است. "
        "طبقه‌بندی قطعی با کتاب تعرفه ایران و ارزیاب گمرک است.",
        size=10,
        color=MUTED,
    )

    # ----- STAGE 2 -----
    _page_break(doc)
    _h(doc, "مرحله ۲ — Import Opportunity Snapshot", 1)
    _add(doc, market.get("imported_statement") or "")
    ev_rows = [
        [e.get("checked_on"), e.get("authority_grade"), e.get("domain"), (e.get("claim") or "")[:90], e.get("relevance"), (e.get("url") or "")[:70]]
        for e in (market.get("imported_evidence") or [])
    ]
    if ev_rows:
        _h(doc, "شواهد پذیرفته‌شده وب (با تاریخ و درجه منبع)", 2)
        _table(doc, ["تاریخ", "درجه", "منبع", "عنوان", "ارتباط", "لینک"], ev_rows)
    _h(doc, "چک‌لیست ریسک‌های مقرراتی — نه حکم قطعی", 2)
    _add(doc, "هر ردیف یک موضوع برای کنترل است. ستون «وضعیت فعلی» نتیجه اولیه ایجنت و ستون «اقدام لازم» کار بعدی کارشناس را نشان می‌دهد.", size=10, color=MUTED)
    _table(
        doc,
        ["موضوع کنترل", "وضعیت فعلی", "این مورد چیست؟", "اقدام لازم"],
        [
            [r["title"], r["level"], r["detail"], r["verification"]]
            for r in (market.get("regulatory_risks") or [])
        ],
    )
    _h(doc, "مراجع رسمی برای کنترل نهایی", 2)
    _add(doc, "درج نام سامانه به معنی انجام‌شدن کنترل نیست. کارشناس باید نتیجه و تاریخ بررسی را ثبت کند.", size=10, color=MUTED)
    _table(
        doc,
        ["سامانه", "چه چیزی کنترل شود", "وضعیت فعلی", "آدرس"],
        [[p["name"], p["check"], p.get("status") or "کنترل دستی الزامی", p["url"]] for p in (market.get("official_portals") or [])],
    )
    for n in market.get("opportunity_notes") or []:
        _add(doc, "• " + n, size=10, color=MUTED)

    # ----- STAGE 3 -----
    _page_break(doc)
    _h(doc, "مرحله ۳ — استراتژی تأمین‌کننده‌یابی و Longlist", 1)
    _h(doc, "Supplier Persona", 2)
    pers = sourcing.get("persona") or {}
    labels = {
        "country_pref": "کشور ترجیحی",
        "company_type": "نوع شرکت",
        "track_record": "سابقه",
        "capacity": "ظرفیت",
        "certificates": "گواهی",
        "moq": "MOQ",
        "export_markets": "بازار صادراتی",
        "terms": "شرایط مطلوب",
    }
    _kv_table(doc, [(labels.get(k, k), v) for k, v in pers.items()])
    _add(
        doc,
        f"پس از فیلتر هویت، ارتباط محصول و حذف تکراری‌ها {len(sourcing.get('longlist') or [])} رکورد از کانال‌های "
        f"{' ، '.join(sourcing.get('channels_used') or [])} در Longlist مانده است. "
        f"{sourcing.get('requirement_status') or ''} جدول کامل و قابل فیلتر در فایل اکسل است.",
    )
    ll = sourcing.get("longlist") or []
    _table(
        doc,
        ["#", "نام", "نام حقوقی", "کشور", "درجه", "تطابق", "کانال", "تماس", "لینک"],
        [
            [
                str(i), s.get("name"), s.get("legal_name") or "تأیید نشده",
                s.get("country"), s.get("candidate_grade"), s.get("product_match"),
                s.get("source_channel"), s.get("contact") or "—",
                (s.get("official_website") or s.get("url") or "")[:65],
            ]
            for i, s in enumerate(ll, 1)
        ],
    )
    _h(doc, "نمونه موارد حذف‌شده (Deduplicate / نامرتبط)", 2)
    _table(
        doc,
        ["دلیل حذف", "عنوان"],
        [[r.get("reason"), (r.get("title") or "")[:90]] for r in (sourcing.get("rejected") or [])[:15]],
    )

    # ----- STAGE 4 -----
    _page_break(doc)
    _h(doc, "مرحله ۴ — ماتریس امتیاز و Top 5", 1)
    _add(doc, scoring.get("model_note") or "")
    _add(doc, scoring.get("top5_status") or "", size=10, color=GOLD)
    _add(
        doc,
        "وزن‌ها: تطابق محصول ۲۰ · حضور دیجیتال ۱۵ · شواهد صادرات ۱۵ · گواهی ۱۰ · "
        "شفافیت هویتی ۱۰ · توان تولید ۱۰ · شرایط تجاری ۱۰ · پاسخ‌گویی واقعی ۱۰. "
        "تا پیش از دریافت پاسخ RFQ، امتیاز پاسخ‌گویی صفر/N/A است.",
        size=10,
    )
    crit = scoring.get("criteria") or []
    headers = ["رتبه", "نام", "کشور", "جمع"] + [c["title"] for c in crit]
    score_rows = []
    for i, s in enumerate(scoring.get("scored") or [], 1):
        scores = s.get("scores") or {}
        score_rows.append(
            [str(i), s.get("name"), s.get("country"), s.get("total")]
            + [scores.get(c["id"], "—") for c in crit]
        )
    _table(doc, headers, score_rows[:22])

    _h(doc, "دلیل امتیاز پنج گزینه برتر", 2)
    for i, s in enumerate(scoring.get("top5") or [], 1):
        _add(doc, f"{i}. {s.get('name')}  —  {s.get('total')}/100  —  {s.get('country')}", bold=True, color=NAVY)
        reasons = s.get("reasons") or {}
        rrows = []
        for c in crit:
            rrows.append([c["title"], f"{(s.get('scores') or {}).get(c['id'], '—')}/{c['max']}", reasons.get(c["id"], "")])
        _table(doc, ["معیار", "نمره", "دلیل ثبت‌شده"], rrows)

    # ----- STAGE 5 -----
    _page_break(doc)
    _h(doc, "مرحله ۵ — کارت‌های Due Diligence", 1)
    _add(
        doc,
        "گواهی فقط در صورت مدرک یا مرجع قابل بررسی معتبر تلقی شده. در این مرحله هیچ گواهی‌ای verify نشده است.",
        size=10,
        color=GOLD,
    )
    for card in cards:
        _h(doc, f"{card.get('name')}   ({card.get('total')}/100)", 2)
        _kv_table(
            doc,
            [
                ("درجه استناد", card.get("citation_grade") or "—"),
                ("وضعیت RFQ", "آماده بازبینی دستی" if card.get("rfq_eligible") else "توقف تا تأیید هویت"),
                ("نام حقوقی اعلام‌شده", card.get("legal_name")),
                ("کشور", card.get("country")),
                ("وب‌سایت رسمی", card.get("official_website") or "تأیید نشد"),
                ("پروفایل کشف", card.get("profile_url")),
                ("ایمیل", card.get("email")),
                ("تلفن", card.get("phone")),
                ("سال تأسیس", card.get("year_founded")),
                ("آدرس", card.get("address")),
                ("اطلاعات ثبتی", card.get("registry")),
                ("گواهی ادعایی", ", ".join(card.get("certs_claimed") or []) or "—"),
            ],
        )
        flags = []
        for g in card.get("green_flags") or []:
            flags.append(["Green", g])
        for r in card.get("red_flags") or []:
            flags.append(["Red", r])
        if flags:
            _table(doc, ["نوع", "پرچم"], flags)
        if card.get("contradictions"):
            _add(doc, "تناقض‌ها:", bold=True)
            for c in card["contradictions"]:
                _add(doc, "• " + c, size=10)

    # ----- STAGE 6 -----
    _page_break(doc)
    _h(doc, "مرحله ۶ — طراحی و شخصی‌سازی ایمیل RFQ", 1)
    _h(doc, "۱) پرامپت‌های استفاده‌شده", 2)
    for key, val in (rfq.get("prompts") or {}).items():
        _add(doc, key, bold=True, rtl=False, color=BLUE)
        _add(doc, val, rtl=False, size=10)

    _h(doc, "۲) نسخه اولیه ایمیل", 2)
    _add(doc, rfq.get("initial_email") or "", rtl=False, size=10)

    _h(doc, "۳) اطلاعات استفاده‌شده برای Personalization", 2)
    prow = []
    for item in rfq.get("personalized") or []:
        f = item.get("personalization_facts") or {}
        prow.append(
            [
                item.get("supplier"),
                f.get("country"),
                f.get("website_or_profile") or "—",
                f.get("citation_grade") or "—",
                " | ".join(f.get("verified_facts_used") or []) or "هیچ fact تأییدشده‌ای استفاده نشد",
                item.get("send_status") or "—",
            ]
        )
    _table(doc, ["شرکت", "کشور", "وب/پروفایل", "درجه استناد", "Factهای استفاده‌شده", "وضعیت ارسال"], prow)

    _h(doc, "۴) سؤال‌های فنی", 2)
    _table(doc, ["#", "سؤال"], [[str(i), q] for i, q in enumerate(rfq.get("technical_questions") or [], 1)], rtl=False)
    _h(doc, "۵) سؤال‌های اعتبارسنجی", 2)
    _table(doc, ["#", "سؤال"], [[str(i), q] for i, q in enumerate(rfq.get("dd_questions") or [], 1)], rtl=False)

    _h(doc, "۶) نسخه نهایی RFQ هر تأمین‌کننده", 2)
    _add(doc, "متن کامل هر ایمیل در فایل جداگانه RFQ هم آمده تا قابل کپی و ارسال باشد.", size=10, color=MUTED)
    for item in rfq.get("personalized") or []:
        _add(doc, item.get("supplier") or "", bold=True, color=NAVY)
        _add(doc, item.get("final_email") or "", rtl=False, size=10)

    _h(doc, "۷) بهبودهای AI نسبت به پیش‌نویس", 2)
    _table(doc, ["#", "بهبود"], [[str(i), x] for i, x in enumerate(rfq.get("improvements") or [], 1)])

    # ----- STAGE 7 -----
    _page_break(doc)
    _h(doc, "مرحله ۷ — مقایسه و انتخاب نهایی", 1)
    _add(doc, decision.get("recommendation_status_fa") or "", bold=True, color=RED)
    cmp_rows = []
    for row in decision.get("comparison") or []:
        role = ""
        if row.get("name") == decision.get("first_choice"):
            role = "گزینه اول"
        elif row.get("name") == decision.get("second_choice"):
            role = "گزینه دوم"
        cmp_rows.append(
            [
                role,
                row.get("name"),
                row.get("total"),
                row.get("country"),
                " | ".join(row.get("strengths") or [])[:160],
                " | ".join(row.get("weaknesses") or [])[:160],
            ]
        )
    _table(doc, ["نقش", "نام", "امتیاز", "کشور", "نقاط قوت", "نقاط ضعف"], cmp_rows)

    _h(doc, "موارد باز برای بررسی کارشناس", 2)
    _table(doc, ["#", "اقدام"], [[str(i), o] for i, o in enumerate(decision.get("open_items") or [], 1)])

    # ----- SOURCES -----
    _h(doc, "Source Log (نمونه در ورد — کامل در اکسل)", 1)
    _add(doc, f"مجموع منابع پذیرفته‌شده و deduplicate‌شده: {len(sources)}. نتیجه خام نامرتبط وارد این شمارش نشده است.")
    src_rows = [
        [s.get("checked_on"), s.get("used_for"), s.get("authority_grade"), s.get("source_type"), (s.get("claim") or "")[:70], s.get("domain"), (s.get("url") or "")[:55]]
        for s in sources[:40]
    ]
    if src_rows:
        _table(doc, ["تاریخ", "کاربرد", "درجه", "نوع", "ادعا", "دامنه", "لینک"], src_rows)

    _h(doc, "محدودیت‌های این پرونده", 1)
    _add(
        doc,
        "صفحات نیازمند ورود، برخی استورفرانت‌های B2B و سامانه‌های رسمی ایران ممکن است از این محیط "
        "قابل واکشی نباشند. امتیاز پاسخ‌گویی قبل از مکاتبه واقعی عمداً صفر/N/A است. "
        "اگر ۲۰ Supplier یا دو گزینه قابل دفاع پیدا نشود، گزارش کمبود می‌دهد و انتخاب ساختگی نمی‌سازد. "
        "کارشناس باید ثبت سفارش، استاندارد، مجوز و تعرفه را در تاریخ تحویل از منبع رسمی کنترل کند.",
        size=10,
    )

    doc.core_properties.title = m.get("project_title") or "پرونده تصمیم‌گیری واردات"
    doc.core_properties.subject = "تحلیل واردات، تأمین‌کننده‌یابی و Due Diligence"
    doc.core_properties.author = m.get("developer_en") or "Setayesh Jafari"
    doc.core_properties.comments = "Generated by TejaratYar — Evidence-first Trade Decision System"
    path = out_dir / f"{file_stems(dossier)}_گزارش_تصمیم‌گیری.docx"
    doc.save(path)
    return path


def build_rfq_docx(dossier: dict, out_dir: Path) -> Path:
    doc = Document()
    rfq = dossier["rfq"]
    m = dossier["meta"]
    _header_footer(doc, m["owner_fa"], m["product_fa"])
    _add(doc, f"بسته ایمیل RFQ — {m['owner_en']} — {m['product_en']}", size=18, bold=True, color=NAVY, rtl=False)
    _add(doc, "Each draft is supplier-specific. Items marked DRAFT must not be sent until legal identity and contact details are verified.", rtl=False, size=10)
    for item in rfq.get("personalized") or []:
        _page_break(doc)
        _add(doc, item.get("supplier") or "", size=16, bold=True, color=NAVY, rtl=False)
        facts = item.get("personalization_facts") or {}
        _kv_table(
            doc,
            [
                ("Company", facts.get("company")),
                ("Country", facts.get("country")),
                ("Profile", facts.get("website_or_profile")),
                ("Citation grade", facts.get("citation_grade") or "—"),
                ("Verified public facts used", " | ".join(facts.get("verified_facts_used") or []) or "None"),
                ("Send status", item.get("send_status") or "manual review required"),
            ],
        )
        _add(doc, item.get("final_email") or "", rtl=False, size=11)
    doc.core_properties.title = f"RFQ Package — {m['product_en']}"
    doc.core_properties.subject = "Supplier-specific RFQ drafts"
    doc.core_properties.author = m.get("developer_en") or "Setayesh Jafari"
    path = out_dir / f"{file_stems(dossier)}_RFQ.docx"
    doc.save(path)
    return path


def build_prompts_txt(dossier: dict, out_dir: Path) -> Path:
    rfq = dossier["rfq"]
    lines = [
        f"پرامپت‌ها و دستورهای اصلی پرونده — {dossier['meta']['owner_fa']} — {dossier['meta']['product_fa']}",
        f"تاریخ: {dossier['meta']['generated_on']}",
        "یادداشت روش: موتور فعلی deterministic + public web search است؛ این دستورها قواعد ایجنت و پرامپت‌های قابل استفاده برای بازبینی AI هستند.",
        "",
        "=== قواعد مراحل ۱ تا ۷ ===",
    ]
    for item in dossier.get("prompt_log") or []:
        lines += ["", f"[{item.get('stage')} — {item.get('purpose')}]", item.get("prompt") or ""]
    lines += ["", "=== مرحله ۶ — تولید و بهبود RFQ ==="]
    for k, v in (rfq.get("prompts") or {}).items():
        lines += ["", f"[{k}]", v]
    lines += [
        "",
        "=== عبارات جستجوی مرحله ۱ ===",
        *[f"- {p}" for p in dossier["brief"].get("search_phrases") or []],
        "",
        "=== پرس‌وجوهای تأمین‌کننده‌یابی ===",
        *[f"- {q}" for q in (dossier["sourcing"].get("queries_used") or [])],
    ]
    path = out_dir / f"{file_stems(dossier)}_روش‌شناسی_و_پرامپت‌ها.txt"
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def build_xlsx(dossier: dict, out_dir: Path) -> Path:
    wb = Workbook()
    wb.properties.creator = dossier.get("meta", {}).get("developer_en") or "Setayesh Jafari"
    wb.properties.title = dossier.get("meta", {}).get("project_title") or "TejaratYar Import Decision Workbook"
    wb.properties.subject = "Supplier sourcing, scoring, due diligence and landed-cost workbook"
    thin = Border(
        left=Side(style="thin", color="C5D0DC"),
        right=Side(style="thin", color="C5D0DC"),
        top=Side(style="thin", color="C5D0DC"),
        bottom=Side(style="thin", color="C5D0DC"),
    )
    head_fill = PatternFill("solid", fgColor="0B4F3A")
    head_font = Font(color="FFFFFF", bold=True, name="Calibri", size=11)
    zebra = PatternFill("solid", fgColor="F4F7FB")
    wrap = Alignment(wrap_text=True, vertical="top", readingOrder=2)
    link_font = Font(name="Calibri", size=10, color="0F9488", underline="single")

    def sheet(title, headers, rows, widths=None, url_cols=None):
        ws = wb.create_sheet(title)
        ws.sheet_view.rightToLeft = True
        ws.freeze_panes = "A2"
        url_cols = url_cols or []
        for col, h in enumerate(headers, 1):
            cell = ws.cell(1, col, h)
            cell.fill = head_fill
            cell.font = head_font
            cell.alignment = Alignment(horizontal="center", wrap_text=True, readingOrder=2)
            cell.border = thin
        for r_i, row in enumerate(rows, 2):
            for c_i, val in enumerate(row, 1):
                cell = ws.cell(r_i, c_i, _safe_excel_value(val))
                cell.alignment = wrap
                cell.border = thin
                cell.font = Font(name="Calibri", size=10)
                if r_i % 2 == 0:
                    cell.fill = zebra
                if c_i in url_cols and isinstance(val, str) and val.startswith("http"):
                    cell.hyperlink = val
                    cell.font = link_font
        if widths:
            for i, w in enumerate(widths, 1):
                ws.column_dimensions[get_column_letter(i)].width = w
        else:
            for i in range(1, len(headers) + 1):
                ws.column_dimensions[get_column_letter(i)].width = 22
        ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{max(1, len(rows)+1)}"
        ws.row_dimensions[1].height = 24
        return ws

    m = dossier["meta"]
    cover = wb.active
    cover.title = "00_راهنما"
    cover.sheet_view.rightToLeft = True
    cover["A1"] = "تجارت‌یار — پرونده حرفه‌ای تصمیم‌گیری واردات"
    cover["A1"].font = Font(name="Calibri", size=18, bold=True, color="0B4F3A")
    guide = [
        ("عنوان پرونده", m.get("project_title")),
        ("تهیه‌کننده گزارش", m["owner_fa"]),
        ("سازمان / مجموعه", m.get("organization") or "—"),
        ("هدف پرونده", m.get("report_purpose") or "—"),
        ("طراحی و توسعه", m.get("developer_fa") or "ستایش جعفری"),
        ("محصول", f"{m['product_fa']} / {m['product_en']}"),
        ("تاریخ", m["generated_on"]),
        ("وضعیت تصمیم", dossier["decision"].get("recommendation_status_fa")),
        ("گزینه اول", dossier["decision"].get("first_choice") or "انتخاب نشد"),
        ("گزینه دوم", dossier["decision"].get("second_choice") or "انتخاب نشد"),
        ("HS کاندید", dossier["brief"].get("hs_primary") or "تأیید نشد"),
        ("QA", (dossier.get("quality_assurance") or {}).get("status")),
        ("تعداد Longlist", len(dossier["sourcing"].get("longlist") or [])),
        ("", ""),
        ("برگه", "محتوا"),
        ("01_Longlist", "همه تأمین‌کنندگان پس از حذف تکراری — فیلتر کانال/کشور"),
        ("02_ماتریس_امتیاز", "نمره همه معیارها برای کل لیست"),
        ("03_دلایل_امتیاز", "دلیل متنی هر معیار برای Top 5"),
        ("04_Top5", "پنج گزینه برتر"),
        ("05_Due_Diligence", "کارت اعتبارسنجی"),
        ("06_مقایسه_انتخاب", "قوت/ضعف و نقش اول/دوم"),
        ("07_Source_Log", "همه لینک‌ها با تاریخ بررسی و پرس‌وجو"),
        ("08_حذف_شده‌ها", "موارد تکراری یا نامرتبط با دلیل"),
    ]
    for i, (k, v) in enumerate(guide, 3):
        cover[f"A{i}"] = k
        cover[f"B{i}"] = v
        cover[f"A{i}"].font = Font(bold=True, color="0F9488")
        cover[f"B{i}"].alignment = Alignment(wrap_text=True)
    cover.column_dimensions["A"].width = 24
    cover.column_dimensions["B"].width = 78

    long_rows = []
    for i, s in enumerate(dossier["sourcing"].get("longlist") or [], 1):
        long_rows.append(
            [
                i,
                s.get("name"), s.get("legal_name") or "",
                s.get("country"), s.get("candidate_grade"), s.get("identity_status"),
                s.get("entity_confidence"), s.get("product_match"),
                s.get("company_type"), s.get("related_product"),
                s.get("contact") or "", s.get("source_channel"), s.get("source_tool") or "",
                s.get("official_website") or "", s.get("url"),
                s.get("product_evidence"), s.get("checked_on"),
            ]
        )
    sheet(
        "01_Longlist",
        ["#", "نام", "نام حقوقی", "کشور", "درجه کاندید", "وضعیت هویت", "اطمینان هویت", "تطابق محصول", "نوع شرکت", "محصول مرتبط", "تماس", "کانال کشف", "ابزار واقعی", "وب‌سایت رسمی", "URL کشف", "شاهد محصول", "تاریخ بررسی"],
        long_rows,
        [5, 28, 30, 14, 12, 24, 12, 12, 28, 22, 28, 18, 24, 42, 42, 55, 14],
        url_cols=[14, 15],
    )

    crit = dossier["scoring"].get("criteria") or []
    headers = ["#", "نام", "کشور", "جمع پیش از RFQ", "واجد Top 5", "دلیل حذف"] + [c["title"] for c in crit] + ["URL"]
    score_rows = []
    for i, s in enumerate(dossier["scoring"].get("scored") or [], 1):
        scores = s.get("scores") or {}
        score_rows.append(
            [i, s.get("name"), s.get("country"), s.get("total"), "بله" if s.get("eligible_for_top5") else "خیر", " | ".join(s.get("disqualify_reasons") or [])]
            + [scores.get(c["id"]) for c in crit]
            + [s.get("url")]
        )
    ws_sc = sheet(
        "02_ماتریس_امتیاز",
        headers,
        score_rows,
        [5, 30, 12, 12, 12, 40] + [14] * len(crit) + [40],
        url_cols=[len(headers)],
    )
    if score_rows:
        ws_sc.conditional_formatting.add(
            f"D2:D{len(score_rows)+1}",
            ColorScaleRule(start_type="min", start_color="FDECEA", mid_type="percentile", mid_value=50, mid_color="FFF4DD", end_type="max", end_color="E5F6EE"),
        )

    reason_rows = []
    for s in dossier["scoring"].get("scored") or []:
        reasons = s.get("reasons") or {}
        scores = s.get("scores") or {}
        for c in crit:
            reason_rows.append([s.get("name"), c["title"], scores.get(c["id"]), c["max"], reasons.get(c["id"], "")])
    sheet(
        "03_دلایل_امتیاز",
        ["تأمین‌کننده", "معیار", "نمره", "سقف", "دلیل ثبت‌شده توسط کارشناس/ایجنت"],
        reason_rows,
        [30, 22, 10, 8, 80],
    )

    top_rows = []
    for i, s in enumerate(dossier["scoring"].get("top5") or [], 1):
        top_rows.append([i, s.get("name"), s.get("total"), s.get("country"), s.get("contact") or "", s.get("url")])
    sheet("04_Top5", ["رتبه", "نام", "امتیاز", "کشور", "تماس", "URL"], top_rows, [8, 32, 10, 14, 28, 46], url_cols=[6])

    dd_rows = []
    for c in dossier.get("cards") or []:
        dd_rows.append(
            [
                c.get("name"), c.get("citation_grade"), "بله" if c.get("rfq_eligible") else "خیر",
                c.get("legal_name"),
                c.get("official_website") or "",
                c.get("profile_url"),
                c.get("country"),
                c.get("address"),
                c.get("phone"),
                c.get("email"),
                c.get("year_founded"),
                c.get("registry"),
                ", ".join(c.get("certs_claimed") or []),
                " | ".join(c.get("green_flags") or []),
                " | ".join(c.get("red_flags") or []),
                " | ".join(c.get("contradictions") or []),
                c.get("total"),
            ]
        )
    sheet(
        "05_Due_Diligence",
        ["نام", "درجه استناد", "واجد RFQ", "نام حقوقی", "وب‌سایت رسمی", "پروفایل", "کشور", "آدرس", "تلفن", "ایمیل", "تأسیس", "ثبتی", "گواهی ادعایی", "Green", "Red", "تناقض", "امتیاز"],
        dd_rows,
        [26, 20, 12, 28, 34, 34, 12, 30, 18, 28, 10, 30, 24, 42, 42, 30, 10],
        url_cols=[5, 6],
    )

    cmp_rows = []
    for row in dossier["decision"].get("comparison") or []:
        role = ""
        if row.get("name") == dossier["decision"].get("first_choice"):
            role = "اول"
        elif row.get("name") == dossier["decision"].get("second_choice"):
            role = "دوم"
        cmp_rows.append(
            [
                role,
                row.get("name"),
                row.get("total"),
                row.get("country"),
                " | ".join(row.get("strengths") or []),
                " | ".join(row.get("weaknesses") or []),
            ]
        )
    sheet("06_مقایسه_انتخاب", ["نقش", "نام", "امتیاز", "کشور", "قوت", "ضعف"], cmp_rows, [10, 28, 10, 14, 55, 55])

    src_rows = [
        [s.get("sid"), s.get("checked_on"), s.get("used_for"), s.get("claim"), s.get("authority_grade"), s.get("source_type"), s.get("relevance"), s.get("evidence_status"), s.get("title"), s.get("domain"), s.get("url"), s.get("query"), s.get("snippet")]
        for s in dossier.get("sources") or []
    ]
    sheet(
        "07_Source_Log",
        ["Evidence ID", "تاریخ", "کاربرد", "ادعای پشتیبانی‌شده", "درجه", "نوع منبع", "ارتباط", "وضعیت", "عنوان", "دامنه", "URL", "پرس‌وجو", "شاهد"],
        src_rows,
        [16, 14, 28, 45, 10, 22, 10, 14, 38, 22, 46, 40, 55],
        url_cols=[11],
    )

    rej = [[r.get("reason"), r.get("title"), r.get("url") or ""] for r in (dossier["sourcing"].get("rejected") or [])]
    sheet("08_حذف_شده‌ها", ["دلیل حذف", "عنوان", "URL"], rej, [36, 50, 46], url_cols=[3])

    trows = []
    for r in dossier.get("tool_log") or []:
        trows.append(
            [
                r.get("stage"),
                r.get("tool"),
                r.get("role"),
                r.get("how") or r.get("method"),
                r.get("hits"),
                " | ".join(r.get("queries") or []),
            ]
        )
    sheet(
        "09_ابزارها",
        ["مرحله", "ابزار", "نقش", "چگونه استفاده شد", "تعداد نتیجه", "پرس‌وجوها"],
        trows,
        [14, 22, 36, 48, 12, 70],
    )

    qa_rows = [["قبول" if q.get("passed") else "ناقص/رد", q.get("check"), q.get("detail")] for q in (dossier.get("quality_assurance") or {}).get("checks") or []]
    sheet("10_QA", ["وضعیت", "کنترل کیفیت", "جزئیات"], qa_rows, [14, 38, 75])

    response_rows = []
    for s in dossier["scoring"].get("top5") or []:
        response_rows.append([s.get("name"), "ارسال نشده", "", "", "", "", "", "", "پس از پاسخ واقعی تکمیل و امتیاز Response Quality بازنگری شود"])
    sheet(
        "11_پاسخ_RFQ",
        ["تأمین‌کننده", "وضعیت ارسال", "تاریخ ارسال", "تاریخ پاسخ", "زمان پاسخ (ساعت)", "Quote کامل؟", "اسناد کامل؟", "نمره پاسخ‌گویی /10", "یادداشت"],
        response_rows,
        [30, 16, 14, 14, 16, 14, 14, 18, 65],
    )

    # Landed-cost scenario template. Blank inputs are deliberate and must not be
    # mistaken for quotes or official tariff values.
    lc = wb.create_sheet("12_Landed_Cost")
    lc.sheet_view.rightToLeft = True
    lc["A1"] = "قالب سناریوی Landed Cost — فقط با Quote و نرخ‌های رسمی تکمیل شود"
    lc["A1"].font = Font(name="Calibri", size=14, bold=True, color="0B4F3A")
    lc.merge_cells("A1:D1")
    lc_rows = [
        ("تأمین‌کننده", "", "INPUT REQUIRED", "نام گزینه و منبع Quote"),
        ("ارز", "", "INPUT REQUIRED", "مثلاً USD/EUR"),
        ("مقدار", "", "INPUT REQUIRED", "واحد باید با Unit Price یکسان باشد"),
        ("قیمت واحد FOB/EXW", "", "INPUT REQUIRED", "شماره و تاریخ Quote"),
        ("ارزش کالا", "=IFERROR(B5*B6,0)", "FORMULA", ""),
        ("حمل داخلی/هزینه مبدأ", "", "INPUT REQUIRED", ""),
        ("کرایه بین‌المللی", "", "INPUT REQUIRED", ""),
        ("بیمه", "", "INPUT REQUIRED", ""),
        ("CIF/Customs value scenario", "=SUM(B7:B10)", "FORMULA", "مبنای قطعی ارزش گمرکی باید رسمی کنترل شود"),
        ("نرخ حقوق ورودی", "", "INPUT REQUIRED", "به‌صورت درصد Excel؛ از منبع رسمی"),
        ("حقوق ورودی", "=B11*B12", "FORMULA", ""),
        ("نرخ مالیات/عوارض", "", "INPUT REQUIRED", "به‌صورت درصد Excel؛ از منبع رسمی"),
        ("مبنای مالیات سناریویی", "=B11+B13", "FORMULA", "قواعد رسمی همان سال کنترل شود"),
        ("مالیات/عوارض", "=B15*B14", "FORMULA", ""),
        ("هزینه‌های محلی/بازرسی/ترخیص", "", "INPUT REQUIRED", ""),
        ("Landed Cost کل", "=B11+B13+B16+B17", "FORMULA", ""),
        ("Landed Cost هر واحد", "=IFERROR(B18/B5,0)", "FORMULA", ""),
    ]
    for col, title in enumerate(["آیتم", "مقدار", "وضعیت", "منبع/یادداشت"], 1):
        c = lc.cell(2, col, title); c.fill = head_fill; c.font = head_font; c.alignment = wrap; c.border = thin
    for r_idx, row in enumerate(lc_rows, 3):
        for c_idx, value in enumerate(row, 1):
            c = lc.cell(r_idx, c_idx, value); c.alignment = wrap; c.border = thin
            if r_idx % 2 == 0: c.fill = zebra
    for col, width in zip("ABCD", [34, 20, 18, 65]): lc.column_dimensions[col].width = width
    lc.freeze_panes = "A3"

    path = out_dir / f"{file_stems(dossier)}_تحلیل_تأمین‌کنندگان.xlsx"
    wb.save(path)
    return path


def export_all(dossier: dict, out_dir: Path) -> dict:
    out_dir.mkdir(parents=True, exist_ok=True)
    files = {
        "report": build_docx(dossier, out_dir),
        "excel": build_xlsx(dossier, out_dir),
        "rfq": build_rfq_docx(dossier, out_dir),
        "prompts": build_prompts_txt(dossier, out_dir),
    }
    return {k: str(v.name) for k, v in files.items()}
